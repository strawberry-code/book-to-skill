"""Extraction strategies (adapters) behind a single ``Extractor`` Protocol.

Each concrete extractor wraps one library or external tool and exposes the same
three-member interface, so the pipeline can try them in order without knowing
which is which. Heavy/optional imports stay lazy inside ``extract`` so importing
this module is cheap and never fails on a missing dependency; a missing library
simply makes ``available()`` return ``False``.
"""

from __future__ import annotations

import html
import html.parser
import posixpath
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Final, Protocol, runtime_checkable
from urllib.parse import unquote

from bookextract.types import ExtractionMode, Figure, PageReporter, log_debug

_PDFTOTEXT_TIMEOUT: Final[int] = 120
_EBOOK_CONVERT_TIMEOUT: Final[int] = 300
_TEXT_ENCODINGS: Final[tuple[str, ...]] = ("utf-8-sig", "utf-8", "cp1252", "latin-1")
_ALL_MODES: Final[tuple[ExtractionMode, ...]] = ("technical", "text")


@runtime_checkable
class Extractor(Protocol):
    """One way to turn a document path into plain text."""

    name: str
    modes: tuple[ExtractionMode, ...]

    def available(self) -> bool:
        """Report whether this strategy can run right now.

        A cheap probe that never imports a heavy library or touches the document
        (typically ``shutil.which`` or ``importlib.util.find_spec``).

        Returns:
            ``True`` if the backing tool/library is present.
        """

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        """Extract plain text from a document.

        Args:
            path: Filesystem path to the document.
            reporter: Optional progress callback, advanced once per page by
                strategies that iterate pages (e.g. pypdf). Others ignore it.

        Returns:
            The extracted text, or ``None`` if this strategy could not handle the
            file (so the chain should try the next one).
        """


class _BothModes:
    """Mixin default: a strategy participates in both PDF modes."""

    modes: tuple[ExtractionMode, ...] = _ALL_MODES


# --------------------------------------------------------------------------- #
# Shared I/O helpers
# --------------------------------------------------------------------------- #


def read_text_file(path: str) -> str | None:
    """Read a text file, trying a few common encodings before giving up.

    Args:
        path: Filesystem path to the text file.

    Returns:
        The decoded contents, or ``None`` if the file is unreadable or matches
        none of the candidate encodings.
    """
    for encoding in _TEXT_ENCODINGS:
        try:
            return Path(path).read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            log_debug(f"read_text_file({encoding}) failed: {exc}")
            return None
    return None


def run_text(cmd: list[str], timeout: int) -> str | None:
    """Run a command and capture its stdout as text.

    Args:
        cmd: The command and arguments to execute.
        timeout: Maximum seconds to wait before aborting.

    Returns:
        The captured stdout if the command exits 0 with non-empty output,
        otherwise ``None``.
    """
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout, check=False)
    except (OSError, subprocess.SubprocessError) as exc:
        log_debug(f"{cmd[0]} failed: {exc}")
        return None
    if result.returncode == 0 and result.stdout.strip():
        return result.stdout
    return None


class _HTMLTextExtractor(html.parser.HTMLParser):
    """Minimal HTML -> plain text converter using stdlib only."""

    SKIP_TAGS: Final[frozenset[str]] = frozenset({"script", "style", "head"})
    BREAK_TAGS: Final[frozenset[str]] = frozenset(
        {"p", "br", "h1", "h2", "h3", "h4", "h5", "h6", "li", "div"}
    )

    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: object) -> None:
        if tag in self.SKIP_TAGS:
            self._skip_depth += 1
        if tag in self.BREAK_TAGS:
            self._parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self.SKIP_TAGS and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._skip_depth:
            self._parts.append(data)

    def get_text(self) -> str:
        return html.unescape("".join(self._parts))


def extract_html_content(raw_html: str) -> str:
    """Convert an HTML string to plain text.

    Uses BeautifulSoup when installed (better whitespace handling), otherwise the
    stdlib :class:`_HTMLTextExtractor`.

    Args:
        raw_html: The raw HTML markup.

    Returns:
        The extracted plain text.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        parser = _HTMLTextExtractor()
        parser.feed(raw_html)
        return parser.get_text()
    soup = BeautifulSoup(raw_html, "html.parser")
    for element in soup(["script", "style", "head"]):
        element.decompose()
    return str(soup.get_text(separator="\n"))


def strip_rtf_fallback(raw: str) -> str:
    """Best-effort RTF to text using regex control-word stripping.

    Args:
        raw: The raw RTF document text.

    Returns:
        Plain text with control words, groups, and hex escapes removed.
    """
    raw = re.sub(r"\\'[0-9a-fA-F]{2}", " ", raw)
    raw = re.sub(r"\\par[d]?", "\n", raw)
    raw = re.sub(r"\\tab", "\t", raw)
    raw = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", raw)
    raw = raw.replace("{", "").replace("}", "")
    return html.unescape(raw)


def _pdf_reader_cls() -> type | None:
    """Return a ``PdfReader`` class from pypdf (preferred) or PyPDF2 (legacy)."""
    try:
        from pypdf import PdfReader

        return PdfReader
    except ImportError:
        pass
    try:
        from PyPDF2 import PdfReader  # type: ignore[no-redef]

        return PdfReader
    except ImportError:
        return None


def _resolve_zip_entry(href: str, opf_dir: str, name_set: frozenset[str]) -> str | None:
    """Map an OPF spine href to its real ZIP entry name.

    Spine hrefs are relative to the OPF file's directory, not the ZIP root, and
    may carry a fragment anchor (``chapter1.xhtml#sec2``) or be percent-encoded
    (``ch%201.xhtml``). Resolve against ``opf_dir`` first; if that misses, fall
    back to a basename match before giving up.

    Args:
        href: The raw href from the OPF manifest/spine.
        opf_dir: The OPF file's directory within the archive (may be empty).
        name_set: All entry names present in the ZIP.

    Returns:
        The matching ZIP entry name, or ``None`` if nothing resolves.
    """
    href = unquote(href.split("#", 1)[0])
    if not href:
        return None
    candidate = posixpath.normpath(posixpath.join(opf_dir, href)) if opf_dir else href
    if candidate in name_set:
        return candidate
    if href in name_set:
        return href
    base = posixpath.basename(href)
    for name in name_set:
        if posixpath.basename(name) == base:
            return name
    return None


# --------------------------------------------------------------------------- #
# PDF strategies
# --------------------------------------------------------------------------- #


class PdftotextExtractor(_BothModes):
    name = "pdftotext"

    def available(self) -> bool:
        return shutil.which("pdftotext") is not None

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        return run_text(["pdftotext", "-layout", path, "-"], _PDFTOTEXT_TIMEOUT)


class PypdfExtractor(_BothModes):
    name = "pypdf"

    def available(self) -> bool:
        return _pdf_reader_cls() is not None

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        reader_cls = _pdf_reader_cls()
        if reader_cls is None:
            return None
        try:
            with Path(path).open("rb") as handle:
                return self._read_pages(reader_cls(handle).pages, reporter)
        except (OSError, ValueError) as exc:
            log_debug(f"pypdf extraction failed: {exc}")
            return None

    @staticmethod
    def _read_pages(pages: object, reporter: PageReporter | None) -> str:
        parts = []
        for page in pages:  # type: ignore[attr-defined]
            parts.append(PypdfExtractor._page_text(page))
            if reporter is not None:
                reporter(1)
        return "\n".join(parts)

    @staticmethod
    def _page_text(page: object) -> str:
        try:
            return page.extract_text() or ""  # type: ignore[attr-defined]
        except Exception as exc:  # noqa: BLE001 - third-party page parsing is unpredictable
            log_debug(f"pypdf page extract failed: {exc}")
            return ""


class PdfminerExtractor(_BothModes):
    name = "pdfminer"

    def available(self) -> bool:
        import importlib.util

        return importlib.util.find_spec("pdfminer") is not None

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        try:
            from pdfminer.high_level import extract_text
        except ImportError:
            return None
        try:
            return str(extract_text(path))
        except Exception as exc:  # noqa: BLE001 - pdfminer raises many ad-hoc errors
            log_debug(f"pdfminer extraction failed: {exc}")
            return None


def _figure_page(picture: object) -> int:
    """Physical page of a Docling figure from its first provenance item, else 0."""
    prov = getattr(picture, "prov", None) or []
    if prov:
        return int(getattr(prov[0], "page_no", 0) or 0)
    return 0


def _figure_kind(picture: object) -> str | None:
    """Best-effort figure classification from Docling annotations, else None."""
    for annotation in getattr(picture, "annotations", None) or []:
        classes = getattr(annotation, "predicted_classes", None) or []
        if classes:
            name = getattr(classes[0], "class_name", None)
            if name:
                return str(name)
        kind = getattr(annotation, "kind", None)
        if kind:
            return str(kind)
    return None


def docling_figures(document: object) -> list[Figure]:
    """Capture captioned figures from a Docling document model (#8).

    Only figures with a non-empty caption are kept — the caption is the verbatim,
    citable handle ``figures.md`` summarizes. Best-effort and defensive: any figure
    that raises is skipped, never failing the surrounding extraction.
    """
    figures: list[Figure] = []
    for picture in getattr(document, "pictures", None) or []:
        try:
            caption = picture.caption_text(document).strip()
        except Exception as exc:  # noqa: BLE001 - docling caption resolution is backend-specific
            log_debug(f"docling caption failed: {exc}")
            continue
        if caption:
            figure = Figure(page=_figure_page(picture), caption=caption, kind=_figure_kind(picture))
            figures.append(figure)
    return figures


class DoclingExtractor:
    """Layout-aware extraction. Technical mode only (slow, table/code aware).

    Beyond markdown text it captures captioned figures (#8) from the document model;
    the pipeline pulls them via :meth:`pop_figures` after a winning extraction.
    """

    name = "docling"
    modes: tuple[ExtractionMode, ...] = ("technical",)

    def __init__(self) -> None:
        self._figures: list[Figure] = []

    def available(self) -> bool:
        import importlib.util

        return importlib.util.find_spec("docling") is not None

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        try:
            converter = self._build_converter()
        except ImportError:
            return None
        try:
            document = converter.convert(path).document  # type: ignore[attr-defined]
            markdown = str(document.export_to_markdown())
        except Exception as exc:  # noqa: BLE001 - docling surfaces backend-specific errors
            log_debug(f"docling extraction failed: {exc}")
            return None
        self._figures = docling_figures(document)
        return markdown

    def pop_figures(self) -> list[Figure]:
        """Return figures captured by the last ``extract`` and clear them."""
        figures, self._figures = self._figures, []
        return figures

    @staticmethod
    def _build_converter() -> object:
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import DocumentConverter, PdfFormatOption

        options = PdfPipelineOptions()
        options.do_ocr = False
        options.do_table_structure = True
        return DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=options)}
        )


# --------------------------------------------------------------------------- #
# EPUB strategies
# --------------------------------------------------------------------------- #


class EbooklibExtractor(_BothModes):
    name = "ebooklib"

    def available(self) -> bool:
        import importlib.util

        return all(importlib.util.find_spec(m) for m in ("ebooklib", "bs4"))

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        try:
            import ebooklib
            from ebooklib import epub
        except ImportError:
            return None
        try:
            book = epub.read_epub(path)
            parts = []
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                html_bytes = item.get_content().decode("utf-8", errors="replace")
                parts.append(extract_html_content(html_bytes))
                if reporter is not None:
                    reporter(1)
            return "\n\n".join(parts)
        except Exception as exc:  # noqa: BLE001 - ebooklib raises varied parse errors
            log_debug(f"ebooklib extraction failed: {exc}")
            return None


class ZipfileEpubExtractor(_BothModes):
    """stdlib-only EPUB extractor: unzip, follow the OPF spine, parse HTML."""

    name = "zipfile"

    def available(self) -> bool:
        return True

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        try:
            with zipfile.ZipFile(path) as zf:
                return self._read_archive(zf, reporter)
        except (OSError, zipfile.BadZipFile) as exc:
            log_debug(f"stdlib EPUB extraction failed: {exc}")
            return None

    def _read_archive(self, zf: zipfile.ZipFile, reporter: PageReporter | None) -> str | None:
        names = frozenset(zf.namelist())
        ordered = self._spine_order(zf, names) or self._fallback_order(names)
        if not ordered:
            return None
        body = []
        for name in ordered:
            text = self._read_entry(zf, name)
            if text is not None:
                body.append(text)
            if reporter is not None:
                reporter(1)
        return "\n\n".join(body) if body else None

    @staticmethod
    def _fallback_order(names: frozenset[str]) -> list[str]:
        return sorted(n for n in names if n.endswith((".html", ".xhtml")))

    @staticmethod
    def _spine_order(zf: zipfile.ZipFile, names: frozenset[str]) -> list[str]:
        opf_files = [n for n in names if n.endswith(".opf")]
        if not opf_files:
            return []
        opf_dir = posixpath.dirname(opf_files[0])
        opf_text = zf.read(opf_files[0]).decode("utf-8", errors="replace")
        hrefs = re.findall(r'href=["\']([^"\']+\.(?:xhtml|html))["\']', opf_text)
        resolved = []
        for href in hrefs:
            entry = _resolve_zip_entry(href, opf_dir, names)
            if entry:
                resolved.append(entry)
            else:
                log_debug(f"EPUB spine href unresolved: {href}")
        return resolved

    @staticmethod
    def _read_entry(zf: zipfile.ZipFile, name: str) -> str | None:
        try:
            raw = zf.read(name).decode("utf-8", errors="replace")
        except (KeyError, OSError) as exc:
            log_debug(f"EPUB entry read failed ({name}): {exc}")
            return None
        parser = _HTMLTextExtractor()
        parser.feed(raw)
        return parser.get_text()


# --------------------------------------------------------------------------- #
# DOCX strategies
# --------------------------------------------------------------------------- #


class PythonDocxExtractor(_BothModes):
    name = "python-docx"

    def available(self) -> bool:
        import importlib.util

        return importlib.util.find_spec("docx") is not None

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        try:
            import docx
        except ImportError:
            return None
        try:
            document = docx.Document(path)
        except Exception as exc:  # noqa: BLE001 - python-docx raises package-specific errors
            log_debug(f"python-docx extraction failed: {exc}")
            return None
        return self._render(document)

    @staticmethod
    def _render(document: object) -> str:
        parts = [p.text for p in document.paragraphs if p.text]  # type: ignore[attr-defined]
        for table in document.tables:  # type: ignore[attr-defined]
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)


class ZipfileDocxExtractor(_BothModes):
    name = "zipfile-docx"

    def available(self) -> bool:
        return True

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        import xml.etree.ElementTree as ET

        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        try:
            with zipfile.ZipFile(path) as zf:
                root = ET.fromstring(zf.read("word/document.xml"))
        except (OSError, zipfile.BadZipFile, KeyError, ET.ParseError) as exc:
            log_debug(f"stdlib DOCX extraction failed: {exc}")
            return None
        parts = [
            "".join(t.text for t in para.iter(f"{ns}t") if t.text) for para in root.iter(f"{ns}p")
        ]
        body = [p for p in parts if p]
        return "\n".join(body) if body else None


# --------------------------------------------------------------------------- #
# RTF / HTML / plain-text / Calibre strategies
# --------------------------------------------------------------------------- #


class StriprtfExtractor(_BothModes):
    name = "striprtf"

    def available(self) -> bool:
        import importlib.util

        return importlib.util.find_spec("striprtf") is not None

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        raw = read_text_file(path)
        if raw is None:
            return None
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError:
            return None
        try:
            text = rtf_to_text(raw)
        except Exception as exc:  # noqa: BLE001 - striprtf can raise on malformed input
            log_debug(f"striprtf extraction failed: {exc}")
            return None
        return text if text.strip() else None


class RtfRegexExtractor(_BothModes):
    name = "rtf-regex"

    def available(self) -> bool:
        return True

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        raw = read_text_file(path)
        return strip_rtf_fallback(raw) if raw is not None else None


class HtmlExtractor(_BothModes):
    name = "html-parser"

    def available(self) -> bool:
        return True

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        raw = read_text_file(path)
        return extract_html_content(raw) if raw is not None else None


class PlainTextExtractor(_BothModes):
    name = "plain-text"

    def available(self) -> bool:
        return True

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        return read_text_file(path)


class EbookConvertExtractor(_BothModes):
    """MOBI/AZW/AZW3 via Calibre's ``ebook-convert`` external command."""

    name = "ebook-convert"

    def available(self) -> bool:
        return shutil.which("ebook-convert") is not None

    def extract(self, path: str, reporter: PageReporter | None = None) -> str | None:
        with tempfile.TemporaryDirectory() as tmp:
            out_path = Path(tmp) / "ebook-convert-output.txt"
            try:
                result = subprocess.run(
                    ["ebook-convert", path, str(out_path)],
                    capture_output=True,
                    text=True,
                    timeout=_EBOOK_CONVERT_TIMEOUT,
                    check=False,
                )
            except (OSError, subprocess.SubprocessError) as exc:
                log_debug(f"ebook-convert failed: {exc}")
                return None
            if result.returncode != 0 or not out_path.exists():
                return None
            text = out_path.read_text(encoding="utf-8", errors="replace")
        return text if text.strip() else None
